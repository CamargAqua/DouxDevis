"""Génération du devis DOUX Joaillier au format .docx."""
from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


GREY_BG = "D9D9D9"
DOUX_GOLD = RGBColor(0xC8, 0xA8, 0x55)

# Frais de refus par marque (case-insensitive). Absent = pas de mention.
FRAIS_REFUS: dict[str, int] = {
    "dinh van":       25,
    "breitling":      30,
    "baume & mercier": 30,
    "panerai":        80,
    "zénith":        100,
    "zenith":        100,   # alias sans accent
}


def _set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "000000")
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def _add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False,
             size: int = 10, color: RGBColor | None = None) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        run.text = ""


def _add_header(doc: Document, marque: str) -> None:
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Cm(9)
    header_table.columns[1].width = Cm(9)

    left = header_table.cell(0, 0)
    left.width = Cm(9)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(p, "DOUX", bold=True, size=44, color=RGBColor(0, 0, 0))
    sub = left.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(sub, "JOAILLIER", bold=True, size=12, color=DOUX_GOLD)

    right = header_table.cell(0, 1)
    right.width = Cm(9)
    p2 = right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(p2, marque.upper(), bold=True, size=24)
    doc.add_paragraph()


def _add_client_row(doc: Document, nom: str, sav: str, date_lieu: str) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = (Cm(6), Cm(5.5), Cm(6.5))
    for cell, width, text in zip(table.rows[0].cells, widths, (nom, f"SAV  {sav}", date_lieu)):
        cell.width = width
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        _add_run(p, text, size=11)


def _add_section_title(doc: Document, title: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_bg(cell, GREY_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, title, bold=True, size=11)
    run = p.runs[0]
    run.underline = True


def _add_montre_table(doc: Document, montre: dict[str, Any], photo_bytes: bytes | None,
                      marque: str = "") -> None:
    """Tableau montre redesigné : en-tête sombre (modèle) + corps (état | ref/série)."""
    modele_raw  = (montre.get("modele") or "").upper()
    metal       = (montre.get("metal") or "").upper()
    modele_full = f"{modele_raw} — {metal}" if (modele_raw and metal) else modele_raw
    ref         = montre.get("reference") or ""
    serie       = montre.get("numero_serie") or ""

    etat_lines = montre.get("etat") or []
    if isinstance(etat_lines, str):
        etat_lines = [l.strip() for l in etat_lines.splitlines() if l.strip()]

    has_info = any([modele_raw, ref, serie, etat_lines])

    # ── Cas sans info : ligne unique ─────────────────────────────────────────
    if not has_info:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        default = f"Montre {marque.upper()}" + (f" — {modele_raw}" if modele_raw else "")
        _add_run(tbl.cell(0, 0).paragraphs[0], default, italic=True, size=10)
        return

    # ── Layout : 2 lignes × 2 cols (+ col photo optionnelle) ────────────────
    DARK_HEX = "1A1814"
    col_left  = Cm(12) if not photo_bytes else Cm(7.5)
    col_right = Cm(6)
    col_photo = Cm(4.5)

    if photo_bytes:
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        tbl.autofit = False
        tbl.rows[0].cells[0].width = col_photo
        tbl.rows[0].cells[1].width = col_left
        tbl.rows[0].cells[2].width = col_right
    else:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.autofit = False
        tbl.rows[0].cells[0].width = col_left
        tbl.rows[0].cells[1].width = col_right

    # Photo (col 0, ligne unique)
    if photo_bytes:
        ph = tbl.cell(0, 0)
        ph.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = ph.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pp.add_run()
        try:
            run.add_picture(BytesIO(photo_bytes), width=Cm(4))
        except Exception:
            _add_run(pp, "[photo]", italic=True, size=9)
        off = 1
    else:
        off = 0  # offset col index

    # Ligne unique : gauche = marque + modèle + état | droite = réf + série
    body_l = tbl.cell(0, off)
    body_r = tbl.cell(0, off + 1)
    body_l.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    body_r.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Marque (petit) + Modèle (gras)
    p_marque = body_l.paragraphs[0]
    _add_run(p_marque, marque.upper(), size=8, color=RGBColor(0x88, 0x88, 0x88))
    p_modele = body_l.add_paragraph()
    _add_run(p_modele, modele_full, bold=True, size=11)

    # État si présent
    if etat_lines:
        body_l.add_paragraph()  # espace
        for line in etat_lines:
            p = body_l.add_paragraph()
            _add_run(p, f"• {line.upper()}", size=10)

    # Référence + N° série (une seule fois, à droite)
    first = True
    if ref:
        p = body_r.paragraphs[0]
        first = False
        _add_run(p, "RÉFÉRENCE :", bold=True, size=9)
        p2 = body_r.add_paragraph()
        _add_run(p2, ref, size=10)
    if serie:
        if not first:
            body_r.add_paragraph()  # espace
        p = body_r.paragraphs[0] if first else body_r.add_paragraph()
        first = False
        _add_run(p, "N° DE SÉRIE :", bold=True, size=9)
        p2 = body_r.add_paragraph()
        _add_run(p2, serie, size=10)
    if first:
        _add_run(body_r.paragraphs[0], "—", italic=True, size=10)


def _add_work_table(doc: Document, title: str, lines: list[dict[str, Any]],
                    *, total_label: str | None = None, total_value: float | None = None,
                    intro: str | None = None) -> None:
    COL_DESC  = Cm(13.5)
    COL_PRICE = Cm(4)

    header = doc.add_table(rows=1, cols=2)
    header.style = "Table Grid"
    header.autofit = False
    header.columns[0].width = COL_DESC
    header.columns[1].width = COL_PRICE
    h_left = header.cell(0, 0)
    h_right = header.cell(0, 1)
    h_left.width = COL_DESC
    h_right.width = COL_PRICE
    _set_cell_bg(h_left, GREY_BG)
    _set_cell_bg(h_right, GREY_BG)
    p_left = h_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_left, title, bold=True, size=11)
    p_left.runs[0].underline = True
    p_right = h_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_right, "PRIX TTC EN EUR", bold=True, size=10)
    p_right.runs[0].underline = True

    body_rows = max(len(lines), 1)
    body = doc.add_table(rows=body_rows, cols=2)
    body.style = "Table Grid"
    body.autofit = False
    body.columns[0].width = COL_DESC
    body.columns[1].width = COL_PRICE

    for i, line in enumerate(lines):
        desc_cell = body.cell(i, 0)
        price_cell = body.cell(i, 1)
        desc_cell.width = COL_DESC
        price_cell.width = COL_PRICE
        _clear_paragraph(desc_cell.paragraphs[0])

        description = (line.get("description") or "").strip()
        _SERVICE_KEYWORDS = ("SERVICE", "RÉVISION", "REVISION", "OVERHAUL", "GENERAL SERVICE", "ENTRETIEN")
        is_service = any(kw in (description or "").upper() for kw in _SERVICE_KEYWORDS)
        if i == 0 and intro and is_service:
            p = desc_cell.paragraphs[0]
            _add_run(p, description.upper() if description else "", bold=True, size=10)
            for intro_line in intro.split("\n"):
                if intro_line.strip():
                    extra = desc_cell.add_paragraph()
                    _add_run(extra, intro_line.strip(), italic=True, size=9)
        else:
            p = desc_cell.paragraphs[0]
            _add_run(p, description.upper(), size=10)

        price_p = price_cell.paragraphs[0]

        price_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Utiliser prix_client (calculé par le backend) si disponible, sinon prix partenaire
        prix = line.get("prix_client") if line.get("prix_client") is not None else line.get("prix", 0)
        prix_text = line.get("prix_label") or _format_price(prix)
        _add_run(price_p, prix_text, size=10)

    if not lines:
        empty = body.cell(0, 0)
        _add_run(empty.paragraphs[0], "", size=10)

    if total_label is not None and total_value is not None:
        total = doc.add_table(rows=1, cols=2)
        total.style = "Table Grid"
        total.autofit = False
        total.columns[0].width = COL_DESC
        total.columns[1].width = COL_PRICE
        tl = total.cell(0, 0)
        tr = total.cell(0, 1)
        tl.width = COL_DESC
        tr.width = COL_PRICE
        p = tl.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_run(p, total_label, italic=True, bold=True, size=10)
        p2 = tr.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        try:
            _tv = float(total_value or 0)
            _total_str = f"{_tv:,.2f}".replace(",", " ").replace(".", ",") + " €"
        except (TypeError, ValueError):
            _total_str = "0,00 €"
        _add_run(p2, _total_str, bold=True, size=10)


def _format_price(value: Any) -> str:
    if value is None or value == "":
        return ""
    if isinstance(value, str):
        try:
            value = float(value.replace(",", "."))
        except ValueError:
            return value
    if value == 0:
        return "OFFERT"
    return f"{value:,.2f}".replace(",", " ").replace(".", ",").replace(" ", " ")


def _add_footer_block(doc: Document, delai: str, frais_refus: int | None = None) -> None:
    p = doc.add_paragraph()
    _add_run(p, f"DELAIS APRES ACCORD {delai.upper()} SOUS RESERVE DE DISPONIBILITE DES PIECES",
             size=9)

    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(11)
    table.columns[1].width = Cm(7)

    for row in table.rows:
        for cell in row.cells:
            for edge in ("top", "left", "bottom", "right"):
                tc_pr = cell._tc.get_or_add_tcPr()
                borders = tc_pr.find(qn("w:tcBorders"))
                if borders is None:
                    borders = OxmlElement("w:tcBorders")
                    tc_pr.append(borders)
                b = OxmlElement(f"w:{edge}")
                b.set(qn("w:val"), "nil")
                borders.append(b)

    accord = table.cell(0, 0)
    refus  = table.cell(1, 0)
    _add_run(accord.paragraphs[0], "☐  ACCORD AU DEVIS", size=10)
    _add_run(refus.paragraphs[0],  "☐  REFUS DU DEVIS",  size=10)
    # Frais de refus affichés sous la ligne REFUS dans la même cellule
    if frais_refus is not None:
        p_frais = refus.add_paragraph()
        _add_run(p_frais,
                 f"Merci de noter qu'un refus du devis entraînera des frais de {frais_refus} €.",
                 italic=True, size=9)

    sig_cell = table.cell(0, 1)
    sig_cell.merge(table.cell(1, 1))
    _set_cell_borders(sig_cell)
    _add_run(sig_cell.paragraphs[0], "DATE ET SIGNATURE :", size=10)
    sig_cell.add_paragraph()
    sig_cell.add_paragraph()

    doc.add_paragraph()
    legal = doc.add_paragraph()
    legal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(legal,
             "SARL DOUX Développement au Capital de 15 245 € - R.C. Avignon 65 A 59 – Siret 315 215 442 00023 – APE 4777 Z – TVA",
             size=8)
    legal2 = doc.add_paragraph()
    legal2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(legal2, "Intracommunautaire FR 313 152 15442  ", size=8)
    _add_run(legal2, "sav@douxjoaillier.com", size=8, color=RGBColor(0x05, 0x63, 0xC1))

    # Ligne CGV + QR code côte à côte
    cgv_table = doc.add_table(rows=1, cols=2)
    cgv_table.autofit = False
    cgv_table.columns[0].width = Cm(15)
    cgv_table.columns[1].width = Cm(3)
    for cell in cgv_table.rows[0].cells:
        for edge in ("top", "left", "bottom", "right"):
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn("w:tcBorders"))
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "nil")
            borders.append(b)

    cgv_cell = cgv_table.cell(0, 0)
    cgv_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_cgv = cgv_cell.paragraphs[0]
    p_cgv.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_cgv = p_cgv.add_run(
        "Toute signature du présent devis vaut acceptation des CGV "
        "consultables via le QR code apposé sur ce devis"
    )
    run_cgv.font.size = Pt(7)
    run_cgv.font.underline = True
    run_cgv.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    qr_cell = cgv_table.cell(0, 1)
    qr_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    qr_path = Path(__file__).parent / "static" / "qr_cgv.png"
    if qr_path.exists():
        p_qr = qr_cell.paragraphs[0]
        p_qr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_qr.add_run().add_picture(str(qr_path), width=Cm(2.0))


def build_docx(data: dict[str, Any], photo_bytes: bytes | None = None) -> bytes:
    """Construit le devis DOUX et renvoie le contenu binaire .docx."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    marque = data.get("marque") or "PARTENAIRE"
    client = data.get("client") or {}
    sav = data.get("sav") or {}
    montre = data.get("montre") or {}

    _add_header(doc, marque)

    # Titre "DEVIS DE RÉPARATION"
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titre.paragraph_format.space_before = Pt(4)
    titre.paragraph_format.space_after = Pt(4)
    _add_run(titre, "DEVIS DE RÉPARATION", bold=True, size=14)

    _add_client_row(
        doc,
        nom=(client.get("nom") or "").upper(),
        sav=sav.get("numero", ""),
        date_lieu=f"Le {sav.get('date', '')} à {sav.get('lieu', 'Avignon')}",
    )

    # Phrase d'introduction dynamique (avec nom de la pièce si disponible)
    _m = montre
    intro_p = doc.add_paragraph()
    intro_p.paragraph_format.space_before = Pt(4)
    intro_p.paragraph_format.space_after = Pt(4)
    _add_run(intro_p,
             "Madame, Monsieur,\n"
             "Suite à l'examen de votre montre, veuillez trouver ci-dessous "
             "nos préconisations de remise en état.",
             italic=True, size=10)

    _add_section_title(doc, "INFORMATIONS DE LA MONTRE")
    _add_montre_table(doc, montre, photo_bytes, marque=marque)

    necessaires = data.get("interventions_necessaires") or []
    intro = (data.get("service_complet_description") or "").strip() or None
    total_ttc = data.get("total_ttc")
    if total_ttc in (None, "", 0) and necessaires:
        try:
            # Utiliser prix_client si disponible, sinon prix partenaire brut
            total_ttc = sum(
                float(line.get("prix_client") if line.get("prix_client") is not None else line.get("prix") or 0)
                for line in necessaires
                if (line.get("prix_label") or "") not in ("OFFERT", "INCL")
            )
        except (TypeError, ValueError):
            total_ttc = 0

    _add_work_table(
        doc,
        "TRAVAIL À RÉALISER",
        necessaires,
        total_label="TOTAL TTC EN EURO HORS OPTIONS",
        total_value=total_ttc,
        intro=intro,
    )

    optionnelles = data.get("interventions_optionnelles") or []
    if optionnelles:
        total_opt = sum(
            float(l.get("prix_client") if l.get("prix_client") is not None else l.get("prix") or 0)
            for l in optionnelles
            if l.get("prix_label") not in ("OFFERT", "INCL")
        )
        total_avec_opt = float(total_ttc or 0) + total_opt
        _add_work_table(doc, "TRAVAIL OPTIONNEL", optionnelles,
                        total_label="TOTAL TTC EN EURO OPTIONS INCLUSES",
                        total_value=total_avec_opt)

    frais_refus = FRAIS_REFUS.get((marque or "").lower())
    _add_footer_block(doc, delai=data.get("delai") or "4 à 6 semaines", frais_refus=frais_refus)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
